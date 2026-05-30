import io
import logging
logger = logging.getLogger(__name__)

def extract_text_from_pdf(file):
    try:
        import pdfplumber
    except ImportError:
        return ''
    text_parts = []
    try:
        if hasattr(file, 'read'):
            file.seek(0)
            pdf_file = io.BytesIO(file.read())
        else:
            pdf_file = file
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
    except Exception as e:
        logger.error(f'PDF extraction error: {e}')
        return ''
    return '\n'.join(text_parts).strip()

def extract_text_from_docx(file):
    try:
        from docx import Document
    except ImportError:
        return ''
    text_parts = []
    try:
        if hasattr(file, 'read'):
            file.seek(0)
            doc_file = io.BytesIO(file.read())
        else:
            doc_file = file
        doc = Document(doc_file)
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
    except Exception as e:
        logger.error(f'DOCX extraction error: {e}')
        return ''
    return '\n'.join(text_parts).strip()

def extract_text(file):
    if file is None:
        return ''
    if hasattr(file, 'name'):
        filename = file.name.lower()
    elif isinstance(file, str):
        filename = file.lower()
    else:
        filename = ''
    if filename.endswith('.pdf'):
        return extract_text_from_pdf(file)
    elif filename.endswith('.docx'):
        return extract_text_from_docx(file)
    return ''

def clean_text(text):
    if not text:
        return ''
    import re
    text = text.replace('\x00', '')
    text = re.sub(r'[ \t]+', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()
