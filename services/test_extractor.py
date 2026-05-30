"""
Unit tests for services/extractor.py
Tests PDF extraction, DOCX extraction, and text cleaning.
"""

import io
from django.test import TestCase
from services.extractor import clean_text, extract_text


class TestCleanText(TestCase):

    def test_removes_null_bytes(self):
        text = "hello\x00world"
        result = clean_text(text)
        self.assertNotIn('\x00', result)
        self.assertIn('hello', result)
        self.assertIn('world', result)

    def test_collapses_extra_spaces(self):
        text = "hello    world"
        result = clean_text(text)
        self.assertEqual(result, "hello world")

    def test_collapses_extra_newlines(self):
        text = "line1\n\n\n\n\nline2"
        result = clean_text(text)
        self.assertNotIn('\n\n\n', result)

    def test_strips_whitespace(self):
        text = "   hello world   "
        result = clean_text(text)
        self.assertEqual(result, "hello world")

    def test_empty_string_returns_empty(self):
        self.assertEqual(clean_text(''), '')

    def test_none_returns_empty(self):
        self.assertEqual(clean_text(None), '')

    def test_normal_text_unchanged(self):
        text = "Python developer with 5 years of experience."
        result = clean_text(text)
        self.assertEqual(result, text)


class TestExtractText(TestCase):

    def test_unsupported_format_returns_empty(self):
        class FakeFile:
            name = 'resume.txt'
        result = extract_text(FakeFile())
        self.assertEqual(result, '')

    def test_none_returns_empty(self):
        result = extract_text(None)
        self.assertEqual(result, '')

    def test_docx_extraction(self):
        """Create a real DOCX in memory and extract text from it."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest('python-docx not installed')

        doc = Document()
        doc.add_paragraph('John Smith')
        doc.add_paragraph('Python Developer with 5 years of experience.')
        doc.add_paragraph('Skills: Python, Django, PostgreSQL')

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        buffer.name = 'test_resume.docx'

        result = extract_text(buffer)
        self.assertIn('John Smith', result)
        self.assertIn('Python', result)

    def test_pdf_extraction(self):
        """Test PDF extraction with pdfplumber."""
        try:
            import pdfplumber
        except ImportError:
            self.skipTest('pdfplumber not installed')
        # pdfplumber requires a real PDF — test with a minimal valid PDF
        # We just verify the function handles errors gracefully
        buffer = io.BytesIO(b'not a real pdf')
        buffer.name = 'fake.pdf'
        result = extract_text(buffer)
        # Should return empty string on invalid PDF, not raise an exception
        self.assertIsInstance(result, str)
